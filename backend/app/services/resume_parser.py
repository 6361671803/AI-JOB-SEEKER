import io

import docx
from pypdf import PdfReader

SUPPORTED_EXTENSIONS = {".pdf", ".docx"}


class UnsupportedResumeFormat(ValueError):
    pass


def extract_text(filename: str, content: bytes) -> str:
    ext = filename.lower().rsplit(".", 1)
    ext = f".{ext[-1]}" if len(ext) == 2 else ""

    if ext == ".pdf":
        return _extract_pdf(content)
    if ext == ".docx":
        return _extract_docx(content)
    raise UnsupportedResumeFormat(f"Unsupported resume format: '{ext or filename}'. Only PDF and DOCX are supported.")


def _extract_pdf(content: bytes) -> str:
    reader = PdfReader(io.BytesIO(content))
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n".join(pages).strip()


def _extract_docx(content: bytes) -> str:
    document = docx.Document(io.BytesIO(content))
    paragraphs = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            paragraphs.extend(cell.text for cell in row.cells)
    return "\n".join(paragraphs).strip()
